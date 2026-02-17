import re
import requests
import streamlit as st
from io import BytesIO
from pypdf import PdfReader
import docx

# PDF generation
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch


# ================= PAGE =================
st.set_page_config(page_title="HireMind AI PRO", layout="centered")
st.title("HireMind AI 🧠 PRO")
st.caption("Professional HR Screening Engine (Cloud-ready)")

# ================= MASTER HR PROMPT =================
# NOTE: Structured TEXT (no JSON / no markdown). We enforce stable scoring + clear evidence rules.
BASE_PROMPT = """
You are a Senior HR Talent Acquisition Specialist and Workforce Strategist.

You will screen ONE candidate CV against ONE Job Description.

STRICT RULES (VERY IMPORTANT):
1) Evidence-based ONLY: if a skill/tool/experience is not clearly mentioned in the CV text, treat it as NOT evidenced.
2) No guessing, no assumptions.
3) Output must be clean structured plain text (NOT JSON, NOT markdown).
4) Consistency requirement:
   - Use the scoring rubric below and show the score breakdown table.
   - Your final Total Score MUST equal the sum of the breakdown.
   - If evidence is weak/unclear, score conservatively.
5) IMPORTANT: When you list "Matched Skills" and "Missing/Weak Skills":
   - Only list skills explicitly required in the Job Description.
   - Matched = clearly present in CV (explicit wording or very direct equivalent).
   - Missing/Weak = required by JD but not found OR only vaguely implied.

=============================
1) EXECUTIVE SUMMARY
=============================
- 4–6 lines maximum.
- Mention strongest alignment and biggest concern.

=============================
2) REQUIREMENT-BY-REQUIREMENT MATCH (from JD)
=============================
Create a table-like list of the JD key requirements and mark each one:
- Status: MATCHED / PARTIAL / MISSING
- Evidence: quote short phrase(s) from CV that prove it (max 12 words each).
Example format per requirement:
- Requirement: <...>
  Status: MATCHED
  Evidence: "<short quote>", "<short quote>"

Then:
A) Matched Requirements (bullet list)
B) Missing/Weak Requirements (bullet list)

=============================
3) TOOLS & SYSTEMS FIT
=============================
- Tools mentioned in JD that appear in CV (with evidence quotes).
- Tools mentioned in JD that do NOT appear in CV.

=============================
4) EXPERIENCE & SENIORITY ANALYSIS
=============================
- Extract years of experience from CV if possible. If not explicit, say "Not clearly stated".
- Compare to JD seniority/years.
- Mention role scope and complexity.

=============================
5) PERFORMANCE & IMPACT INDICATORS
=============================
- Any measurable achievements (numbers/metrics). If none, say "No measurable metrics found".

=============================
6) BEHAVIORAL & SOFT SKILLS (evidence-based)
=============================
- List only what is evidenced (teamwork, communication, leadership, stakeholder mgmt, etc.)

=============================
7) RISK ASSESSMENT
=============================
Classify Risk Level: Low / Medium / High
Provide 3–5 concrete risks (evidence-based).

=============================
8) SCORING MODEL (TOTAL 100)
=============================
Score using this rubric:
- Skills Match (0–30)
- Experience Fit (0–20)
- Domain & Tools Fit (0–20)
- Impact & Achievements (0–15)
- Behavioral & Leadership Fit (0–10)
- Risk Adjustment (-5 to +5)

You MUST output a breakdown table exactly like this:

Score Breakdown:
Skills Match: XX/30
Experience Fit: XX/20
Domain & Tools Fit: XX/20
Impact & Achievements: XX/15
Behavioral & Leadership Fit: XX/10
Risk Adjustment: XX/5
-------------------------
Total Score: XX/100

Then:
Hiring Signal: Strong Hire | Hire | Conditional | Not Recommended
Explain in max 5 lines WHY the signal fits the score.

=============================
9) INTERVIEW STRATEGY
=============================
- 5 Technical Questions (focused on missing/partial requirements)
- 3 Behavioral Questions
- 2 Risk Validation Questions

=============================
10) FINAL HR RECOMMENDATION
=============================
- 3–6 lines, direct recommendation and next step.

Job Description:
<<JOB_DESCRIPTION>>

Candidate CV:
<<CANDIDATE_CV>>
"""

# ================= GROQ (Cloud-ready) =================
# This replaces Ollama and works on Streamlit Cloud.
# We keep temperature=0 for stability.
def call_llm(prompt: str, model: str) -> str:
    api_key = st.secrets.get("GROQ_API_KEY", "")
    if not api_key:
        raise RuntimeError("Missing GROQ_API_KEY in Streamlit Secrets.")

    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are a strict, evidence-based HR screening analyst."},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.0,
        "top_p": 1.0,
        # optional: if Groq supports seed on your model, it helps stability.
        # If it errors, remove "seed".
        "seed": 42,
    }

    r = requests.post(url, headers=headers, json=payload, timeout=90)
    # If seed not supported, retry once without it
    if r.status_code == 400 and "seed" in r.text:
        payload.pop("seed", None)
        r = requests.post(url, headers=headers, json=payload, timeout=90)

    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()


# ================= FILE READING =================
def read_pdf(file):
    reader = PdfReader(BytesIO(file.getvalue()))
    text = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text.append(t)
    return "\n".join(text).strip()

def read_docx(file):
    document = docx.Document(BytesIO(file.getvalue()))
    parts = []
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    # tables (CVs sometimes)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()

def read_txt(file):
    return file.getvalue().decode("utf-8", errors="ignore").strip()

def extract_text(file):
    name = file.name.lower()
    if name.endswith(".pdf"):
        return read_pdf(file)
    if name.endswith(".docx"):
        return read_docx(file)
    if name.endswith(".txt"):
        return read_txt(file)
    return ""


# ================= PDF REPORT =================
def _sanitize_for_pdf(line: str) -> str:
    # Basic sanitization for reportlab Paragraph
    if line is None:
        return ""
    # escape minimal xml entities
    line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return line

def generate_pdf(report_text: str) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    # nicer style for body
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10.5,
        leading=14,
        spaceAfter=6,
    )

    title = styles["Title"]
    title.fontSize = 18
    title.leading = 22

    story = []
    story.append(Paragraph("HireMind AI – Professional HR Screening Report", title))
    story.append(Spacer(1, 0.25 * inch))

    for line in report_text.split("\n"):
        line = line.rstrip()
        if line.strip() == "":
            story.append(Spacer(1, 0.12 * inch))
            continue
        story.append(Paragraph(_sanitize_for_pdf(line), body))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ================= UI =================
with st.sidebar:
    st.subheader("⚙️ Model (Groq)")
    model_name = st.text_input("Model", "llama-3.1-8b-instant")
    st.caption("Deployed on Streamlit Cloud using Groq API. Set GROQ_API_KEY in Secrets.")

st.subheader("Upload Job Description")
jd_file = st.file_uploader("Upload JD (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"], key="jd")

st.subheader("Upload Candidate CV")
cv_file = st.file_uploader("Upload CV (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"], key="cv")

with st.expander("Preview extracted text (optional)"):
    if jd_file:
        st.markdown("**JD preview (first 1500 chars):**")
        st.code(extract_text(jd_file)[:1500])
    else:
        st.write("JD not uploaded yet.")
    if cv_file:
        st.markdown("**CV preview (first 1500 chars):**")
        st.code(extract_text(cv_file)[:1500])
    else:
        st.write("CV not uploaded yet.")

if st.button("Analyze Candidate", use_container_width=True):

    if not jd_file or not cv_file:
        st.error("Please upload both JD and CV files.")
        st.stop()

    jd_text = extract_text(jd_file)
    cv_text = extract_text(cv_file)

    if not jd_text.strip():
        st.error("Could not extract text from JD. If PDF is scanned image, you need OCR.")
        st.stop()
    if not cv_text.strip():
        st.error("Could not extract text from CV. If PDF is scanned image, you need OCR.")
        st.stop()

    prompt = BASE_PROMPT.replace("<<JOB_DESCRIPTION>>", jd_text).replace("<<CANDIDATE_CV>>", cv_text)

    with st.spinner("Analyzing candidate professionally..."):
        try:
            report = call_llm(prompt, model_name.strip())
        except Exception as e:
            st.error("AI call failed. Check Secrets and Logs.")
            st.exception(e)
            st.stop()

    st.subheader("Professional Screening Report")
    st.write(report)

    # Download TXT
    st.download_button(
        "Download Report (TXT)",
        data=report,
        file_name="hiremind_report.txt",
        mime="text/plain",
        use_container_width=True
    )

    # Download PDF
    pdf_bytes = generate_pdf(report)
    st.download_button(
        "Download Report (PDF)",
        data=pdf_bytes,
        file_name="hiremind_report.pdf",
        mime="application/pdf",
        use_container_width=True
    )
